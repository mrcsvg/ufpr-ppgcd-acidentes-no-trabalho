"""Preenche o template da ATIVIDADE 3 com os resultados do projeto.

O texto fica em ``conteudo.py``; aqui esta so a mecanica de escrever nas tabelas
e inserir as figuras, preservando a formatacao do template original.

Uso::

    python -m acidentes_trabalho.figuras     # gera as figuras primeiro
    python reports/atividade3/preencher.py

Requer ``python-docx`` (nao e dependencia do pacote: pip install python-docx).
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
import conteudo as C  # noqa: E402

AQUI = Path(__file__).resolve().parent
RELATORIOS = AQUI.parent

ORIGEM = AQUI / "ATIVIDADE_3_template.docx"
FIGURAS = RELATORIOS / "figuras"
DESTINO = RELATORIOS / "ATIVIDADE_3_preenchida.docx"

doc = docx.Document(ORIGEM)


def escrever(celula, texto, negrito=False, tamanho=9):
    """Escreve texto numa celula, preservando a formatacao do paragrafo do template."""
    celula.text = ""
    p = celula.paragraphs[0]
    run = p.add_run(texto)
    run.font.size = Pt(tamanho)
    run.bold = negrito
    return p


def preencher(tabela, linhas, coluna_inicial=0):
    """Preenche a tabela a partir da linha 1 (a 0 e o cabecalho)."""
    for i, valores in enumerate(linhas, start=1):
        if i >= len(tabela.rows):
            raise IndexError(f"tabela tem {len(tabela.rows)} linhas, precisa de {len(linhas)+1}")
        for j, valor in enumerate(valores):
            escrever(tabela.rows[i].cells[coluna_inicial + j], valor,
                     negrito=(j == 0 and coluna_inicial == 0 and len(valores) > 2))


# --- Tabelas cujo rotulo da 1a coluna ja vem preenchido no template ---
for i, (_, valor) in enumerate(C.TABELA_0, start=1):
    escrever(doc.tables[0].rows[i].cells[1], valor)

for i, (_, resultado, evidencia) in enumerate(C.TABELA_1, start=1):
    escrever(doc.tables[1].rows[i].cells[1], resultado)
    escrever(doc.tables[1].rows[i].cells[2], evidencia, tamanho=8.5)

# --- Tabelas totalmente vazias ---
preencher(doc.tables[2], C.TABELA_2)
preencher(doc.tables[3], C.TABELA_3)
# A 1a coluna ja traz a numeracao no template: o achado entra junto do numero.
preencher(doc.tables[5], [
    (f"{i}. {achado}", evidencia, relevancia, limitacao)
    for i, (achado, evidencia, relevancia, limitacao) in enumerate(C.TABELA_5, 1)
])
preencher(doc.tables[7], C.TABELA_7)

# --- Tabelas com pergunta na 1a coluna ja preenchida ---
for i, (_, resposta) in enumerate(C.TABELA_4, start=1):
    escrever(doc.tables[4].rows[i].cells[1], resposta)
for i, (_, resposta) in enumerate(C.TABELA_6, start=1):
    escrever(doc.tables[6].rows[i].cells[1], resposta)


def paragrafo_apos(referencia):
    """Cria um paragrafo vazio logo depois de ``referencia`` e devolve-o."""
    novo = copy.deepcopy(referencia._p)
    referencia._p.addnext(novo)
    p = docx.text.paragraph.Paragraph(novo, referencia._parent)
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    p.paragraph_format.left_indent = 0
    p.paragraph_format.first_line_indent = 0
    if p._p.pPr is not None and p._p.pPr.numPr is not None:
        p._p.pPr.remove(p._p.pPr.numPr)
    return p


def inserir_figura(apos, arquivo, legenda, largura_cm=15.5):
    """Insere imagem + legenda depois do paragrafo ``apos``; devolve o ultimo criado."""
    p_img = paragrafo_apos(apos)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(FIGURAS / arquivo), width=Cm(largura_cm))

    p_leg = paragrafo_apos(p_img)
    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_leg.add_run(legenda)
    run.font.size = Pt(8.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
    return p_leg


def achar_paragrafo(trecho):
    for p in doc.paragraphs:
        if trecho.lower() in p.text.strip().lower():
            return p
    raise LookupError(f"paragrafo nao encontrado: {trecho!r}")


# As figuras da exploracao entram logo antes de "REFINE — Aprofundar".
ancora = achar_paragrafo("Para cada análise, registre brevemente")
cabecalho = paragrafo_apos(ancora)
run = cabecalho.add_run("Figuras da etapa EXPLORE")
run.bold = True
run.font.size = Pt(10)
atual = cabecalho
for arquivo, legenda in C.FIGURAS:
    atual = inserir_figura(atual, arquivo, legenda)

# O diagrama do workflow entra na secao 6.
ancora = achar_paragrafo("O workflow deve apresentar")
inserir_figura(ancora, "workflow.png", C.LEGENDA_WORKFLOW, largura_cm=16.5)

DESTINO.parent.mkdir(parents=True, exist_ok=True)
doc.save(DESTINO)
print(f"gravado: {DESTINO}  ({DESTINO.stat().st_size/1024:.0f} KB)")
